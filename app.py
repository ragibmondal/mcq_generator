import streamlit as st
import time
import pdfkit
import io
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_BREAK

# Define a dictionary mapping languages to font names
LANGUAGE_FONTS = {
    'English': 'Arial',
    'French': 'Times New Roman',
    'Spanish': 'Calibri',
    'German': 'Verdana',
    'Italian': 'Cambria',
    'Bangla': 'SolaimanLipi',
    'Hindi': 'Mangal',
    'Urdu': 'Jameel Noori Nastaleeq',
}

# Set the page configuration
st.set_page_config(page_title="MCQs Generator App",
                   page_icon="🧐",
                   layout="centered",
                   initial_sidebar_state="auto",
                   )

from src.helper import llm_chain
from src.data_util import read_input_file
from src.logger import logging

# Set the page title
st.title(':red[MCQ] :blue[Generator]')
st.caption('                By :orange[Gemini-Flash-1.5] using Langchain 🐦')

with st.sidebar:
    # uploading the input file
    uploaded_file = st.file_uploader("Choose a PDF | Text file",
                                     accept_multiple_files=False,
                                     type=['txt', 'pdf']
                                     )

    # Number of mcq questions user wants
    number = st.number_input("Insert a number",
                             min_value=1,
                             max_value=100,
                             value=5, placeholder="Type a number...")
    # Difficulty level slider
    level = st.select_slider('Select difficulty',
                             options=['Easy', 'Medium', 'Hard'])

    # Language selection
    language = st.selectbox('Select Language', ['English', 'French', 'Spanish', 'German', 'Italian', 'Bangla', 'Hindi', 'Urdu'])

    if uploaded_file and number and level and language:
        data = read_input_file(uploaded_file)
        gen_button = st.button("Generate", key="gen_button")

try:
    if gen_button:
        with st.spinner('Generating Multi Choice Questions...'):
            # Generating the response from the model
            response = llm_chain.run(number=number,
                                     difficulty=level,
                                     text=data,
                                     language=language)
            # print(response)
        logging.info('MCQ are generated')
except NameError:
    pass

try:
    if gen_button and response:
        # write to UI
        message_placeholder = st.empty()
        full_response = ""
        for chunk in response.replace('\n', '  \n').replace('\t', '----'):
            full_response += chunk
            time.sleep(0.005)
            # Add a blinking cursor to simulate typing
            message_placeholder.markdown(full_response + "▌")
        message_placeholder.markdown(full_response)

        # Word file download
        word_data = create_word_file(response, language)
        st.download_button(
            label="Download Word File",
            data=word_data,
            file_name="mcqs.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # PDF file download
        pdf_data = create_pdf(response)
        st.download_button(
            label="Download PDF File",
            data=pdf_data,
            file_name="mcqs.pdf",
            mime="application/pdf"
        )

except NameError:
    pass

def create_word_file(response, language):
    # Create a new Word document
    doc = Document()

    # Set the default font for the document
    font_name = LANGUAGE_FONTS.get(language, 'Arial')
    doc.styles['Normal'].font.name = font_name

    # Split the response into lines
    lines = response.split('\n')

    # Add each line to the document
    for line in lines:
        doc.add_paragraph(line)
        doc.add_paragraph(WD_BREAK.LINE_BREAK)

    # Save the document to a BytesIO object
    file_obj = io.BytesIO()
    doc.save(file_obj)
    file_obj.seek(0)

    return file_obj.getvalue()

def create_pdf(response):
    pdf = pdfkit.from_string(response, False)
    return pdf
